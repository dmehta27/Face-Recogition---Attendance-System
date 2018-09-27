import numpy as np
import cv2
from docx import Document
from docx.shared import Inches
import sqlite3
import datetime

conn = sqlite3.connect('Attendance.db')
cmd = "SELECT * FROM Attendance_April_2017"
cursor = conn.execute(cmd)

mydate = datetime.datetime.now()
curr_month = mydate.strftime("%B")
curr_year = mydate.strftime("%Y")

document = Document()
document.add_heading("Defaulter List for the Month - "+curr_month+" - "+curr_year)
document.add_paragraph("")

i=2
count=0

for row in cursor:
    attendance=0    
    for i in range (2,len(row)):
         if(row[i]=='P'):
            attendance = attendance+1
    
    if(((attendance*100)/(len(row)-2))<75):
        count=count+1
        document.add_paragraph(str(count)+". "+str(row[1])+" - "+str(row[0]))

        
        
conn.close()
document.save("Defaulter Lists/Defaulter List For - "+curr_month+" - "+curr_year+".docx")
print "Document saved successfully"
