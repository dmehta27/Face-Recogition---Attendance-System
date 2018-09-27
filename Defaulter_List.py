import sqlite3
import numpy as np
from docx import Document
from docx.shared import Inches
import datetime

conn = sqlite3.connect('Attendance.db')
cmd = "SELECT name FROM Student_info"
cursor = conn.execute(cmd)

resultSet = None
document = Document()
document.add_heading("Defaulter List")
document.add_paragraph("")
document.add_paragraph("This is the Defaulter List")
document.add_paragraph("")

count=0
for row in cursor:
     count=count+1
     document.add_paragraph(str(count)+") "+str(row[0]))
        
conn.close()
mydate = datetime.datetime.now()

print mydate.strftime("%d-%m-%Y" )

curr_month = mydate.strftime("%B")
curr_year = mydate.strftime("%Y")

document.save("Defaulter Lists/Defaulter List For - "+curr_month+" - "+curr_year+".docx")
print "Document saved successfully"
